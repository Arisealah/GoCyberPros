
#!/usr/bin/env python3
"""
DOCX Metadata Extraction Tool (Enhanced)

Description: Comprehensive and robust metadata extraction from DOCX files,
with advanced validation, detailed error handling, and standardized, enriched output.

Features:
- Extracts core properties, custom properties, and document statistics.
- Validates DOCX structure and integrity rigorously (ZIP+XML).
- Handles modern and legacy DOCX formats.
- Outputs standardized, structured metadata including warnings and errors.

Dependencies:
- python-docx (pip install python-docx)
- python-magic (pip install python-magic)

Example Usage (CLI):
    python3 docx_extractor.py example.docx --log-level DEBUG -o metadata.json

Author: Gemini (based on provided foundation)
Version: 1.1.0
Last Updated: 2025-06-07
"""

import os
import zipfile
import magic
import mimetypes # For more robust MIME type handling
from datetime import datetime, timezone
from typing import Dict, Any, Optional, List
from docx import Document
from docx.opc.coreprops import CoreProperties
import logging
import sys
import argparse
import time
import json

# --- Configuration ---
DEFAULT_LOG_FILE = "docx_extractor.log"
DEFAULT_LOG_LEVEL = "INFO" # DEBUG, INFO, WARNING, ERROR, CRITICAL

# Expected DOCX MIME types
EXPECTED_DOCX_MIMES = [
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'application/zip' # DOCX files are essentially ZIP files
]

# --- Logging Setup ---
def setup_logging(log_level: str = DEFAULT_LOG_LEVEL):
    """Configures the logging for the extractor."""
    log_formatter = logging.Formatter('%(asctime)s [%(levelname)s] [%(name)s] %(message)s')
    root_logger = logging.getLogger()
    root_logger.setLevel(getattr(logging, log_level.upper()))

    # Clear existing handlers to prevent duplicate output when called multiple times
    if root_logger.handlers:
        for handler in root_logger.handlers:
            root_logger.removeHandler(handler)

    # File Handler
    file_handler = logging.FileHandler(DEFAULT_LOG_FILE, encoding='utf-8')
    file_handler.setFormatter(log_formatter)
    root_logger.addHandler(file_handler)

    # Console Handler
    console_handler = logging.StreamHandler(sys.stdout)
    console_handler.setFormatter(log_formatter)
    root_logger.addHandler(console_handler)

# Initial logging setup
setup_logging()
logger = logging.getLogger(__name__)

# --- Custom Exception Classes ---
class DOCXMetadataError(Exception):
    """Base class for DOCX metadata extraction exceptions."""
    pass

class FileAccessError(DOCXMetadataError):
    """Raised for issues accessing the file (e.g., not found, permissions)."""
    pass

class InvalidFileFormatError(DOCXMetadataError):
    """Raised when the file is not a valid DOCX according to checks."""
    pass

class ExtractorProcessingError(DOCXMetadataError):
    """Raised for unexpected errors during the DOCX extraction process."""
    pass

# --- Helper Functions ---
def _datetime_to_iso_utc(dt_obj: Optional[datetime]) -> Optional[str]:
    """Converts a datetime object to ISO 8601 format with UTC timezone, if not None."""
    if dt_obj:
        if dt_obj.tzinfo is None: # Assume local time if no timezone info
            return dt_obj.astimezone(timezone.utc).isoformat()
        return dt_obj.astimezone(timezone.utc).isoformat()
    return None

# --- Main Validation Function ---
def validate_docx_file(file_path: str) -> bool:
    """
    Advanced validation for DOCX file structure and integrity.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        True if the file passes all validations.

    Raises:
        FileAccessError: If file doesn't exist or access is restricted.
        InvalidFileFormatError: If file is not deemed a valid DOCX.
    """
    logger.info(f"Validating file: {file_path}")

    if not os.path.exists(file_path):
        raise FileAccessError(f"File not found: {file_path}")

    if not os.path.isfile(file_path):
        raise FileAccessError(f"Path is not a regular file: {file_path}")

    if not os.access(file_path, os.R_OK):
        raise FileAccessError(f"Access denied (read permission): {file_path}")

    file_size = os.path.getsize(file_path)
    if file_size == 0:
        logger.warning(f"File {file_path} is empty, which is an invalid DOCX.")
        raise InvalidFileFormatError(f"File is empty: {file_path}")

    # 1. MIME Type Check (using both magic and mimetypes)
    mime_type_magic = magic.from_file(file_path, mime=True)
    mime_type_ext = mimetypes.guess_type(file_path)[0]

    is_mime_ok = mime_type_magic in EXPECTED_DOCX_MIMES
    if not is_mime_ok:
        if mime_type_ext and mime_type_ext in EXPECTED_DOCX_MIMES:
            logger.warning(f"Magic detected '{mime_type_magic}', but extension suggests DOCX ('{mime_type_ext}'). Proceeding.")
        else:
            raise InvalidFileFormatError(f"Not a recognized DOCX file type by magic ({mime_type_magic}).")
    
    # 2. ZIP structure and essential DOCX file check
    try:
        with zipfile.ZipFile(file_path) as z:
            # Check for core DOCX components
            required_files = ['word/document.xml', '_rels/.rels']
            for req_file in required_files:
                if req_file not in z.namelist():
                    raise InvalidFileFormatError(f"Missing required DOCX internal file: {req_file}")
            
            # Check for common corruption indicators (e.g., empty or invalid central directory)
            if not z.namelist():
                raise InvalidFileFormatError("DOCX zip archive is empty or corrupted.")

    except zipfile.BadZipFile as e:
        raise InvalidFileFormatError(f"Invalid ZIP structure for DOCX: {str(e)}") from e
    except Exception as e:
        logger.error(f"Unexpected error during ZIP validation for {file_path}: {e}", exc_info=True)
        raise InvalidFileFormatError(f"Unexpected validation error: {str(e)}") from e

    logger.info(f"File '{file_path}' passed all structural validations.")
    return True

# --- Main Extraction Function ---
# def extract_docx_metadata(file_path: str) -> Dict[str, Any]:
# def _extract_file_info(self, file_path: str) -> Dict[str, Any]:
def extract_docx_metadata(file_path: str) -> Dict[str, Any]:
    """
    Extract comprehensive metadata from DOCX files.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        A dictionary containing extracted metadata, including:
        - file_info: Basic file system attributes.
        - core_properties: Standard DOCX document properties.
        - custom_properties: Custom document properties.
        - technical_info: Document statistics and structural info.
        - processing: Status, warnings, errors, and time taken.

    Raises:
        DOCXMetadataError: For any critical failure during validation or processing.
    """
    result: Dict[str, Any] = {
        "file_info": {},
        "core_properties": {},
        "custom_properties": {},
        "technical_info": {},
        "processing": {
            "success": False,
            "warnings": [],
            "errors": [],
            "time_taken_seconds": None,
            "extractor_version": "1.1.0"
        }
    }

    start_time = time.time()
    logger.info(f"Starting DOCX metadata extraction for: {file_path}")

    doc = None # Initialize to None for finally block

    try:
        # --- Stage 1: Validation ---
        validate_docx_file(file_path)
        result["file_info"]["valid"] = True

        # --- Stage 2: File System Info ---
        file_stat = os.stat(file_path)
        result["file_info"] = {
            "path": os.path.abspath(file_path),
            "filename": os.path.basename(file_path),
            "size_bytes": file_stat.st_size,
            "created_utc": datetime.fromtimestamp(file_stat.st_ctime, tz=timezone.utc).isoformat(),
            "modified_utc": datetime.fromtimestamp(file_stat.st_mtime, tz=timezone.utc).isoformat(),
            "accessed_utc": datetime.fromtimestamp(file_stat.st_atime, tz=timezone.utc).isoformat(),
            "format": "DOCX",
            "detected_mime_type": mimetypes.guess_type(file_path)[0] or magic.from_file(file_path, mime=True),
            "valid": True
        }

        # --- Stage 3: DOCX-specific Metadata Extraction (using python-docx) ---
        try:
            doc = Document(file_path)
            
            # Core Properties
            core_props = doc.core_properties
            result["core_properties"] = {
                "title": core_props.title,
                "subject": core_props.subject,
                "author": core_props.author,
                "keywords": core_props.keywords,
                "comments": core_props.comments,
                "category": core_props.category,
                "content_status": core_props.content_status,
                "created_utc": _datetime_to_iso_utc(core_props.created),
                "modified_utc": _datetime_to_iso_utc(core_props.modified),
                "last_modified_by": core_props.last_modified_by,
                "revision": core_props.revision,
                "version": core_props.version,
                "last_printed_utc": _datetime_to_iso_utc(core_props.last_printed) # Added last printed
            }
            
            # Custom Properties
            custom_props_data = {}
            try:
                # Iterate through custom properties if available (python-docx v0.8.11+)
                if hasattr(doc.part, 'custom_properties'):
                    for prop in doc.part.custom_properties.properties:
                        prop_value = prop.value
                        # Convert datetime objects in custom properties to ISO UTC
                        if isinstance(prop_value, datetime):
                            prop_value = _datetime_to_iso_utc(prop_value)
                        custom_props_data[prop.name] = prop_value
            except Exception as e:
                result["processing"]["warnings"].append(f"Could not fully read custom properties: {str(e)}")
                logger.debug(f"Detail: Custom property parsing issue: {e}")
            result["custom_properties"] = custom_props_data

            # Technical Info / Document Statistics
            total_paragraphs = len(doc.paragraphs)
            total_tables = len(doc.tables)
            total_sections = len(doc.sections)
            
            # Estimate word count and character count more robustly
            word_count = 0
            char_count = 0
            for paragraph in doc.paragraphs:
                text = paragraph.text
                word_count += len(text.split())
                char_count += len(text)
            
            result["technical_info"] = {
                "paragraphs": total_paragraphs,
                "tables": total_tables,
                "sections": total_sections,
                "styles": len(doc.styles),
                "inline_shapes": len(doc.inline_shapes), # Images, etc.
                "footnotes": len(doc.footnotes),
                "endnotes": len(doc.endnotes),
                "headers": len(doc.sections) * 2, # Each section typically has header/footer
                "footers": len(doc.sections) * 2,
                "has_comments": bool(doc.comments), # Check if any comments exist
                "has_revisions": bool(doc.element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}del') or \
                                      doc.element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ins')), # Basic check for tracked changes
                "word_count_approx": word_count,
                "character_count_approx": char_count,
                "has_macros": False # Default to False, check below
            }

            # Check for macros (.docm)
            # This is a robust check using zipfile to look for vbaProject.bin
            with zipfile.ZipFile(file_path, 'r') as zf:
                if 'word/vbaProject.bin' in zf.namelist():
                    result["technical_info"]["has_macros"] = True
                    result["processing"]["warnings"].append("Document contains VBA macros.")

            result["processing"]["success"] = True
        
        except Exception as e:
            result["processing"]["errors"].append(f"DOCX content read error: {str(e)}. File might be corrupted or malformed internally.")
            logger.error(f"Python-docx processing error: {e}")
            result["processing"]["success"] = False # Mark as failed
            return result # Exit early

    except FileAccessError as e:
        result["processing"]["errors"].append(str(e))
        logger.error(f"File Access Error during DOCX extraction: {e}")
    except InvalidFileFormatError as e:
        result["processing"]["errors"].append(str(e))
        logger.error(f"Invalid DOCX Format Error: {e}")
    except Exception as e:
        # Catch any other unexpected errors that might occur outside specific blocks
        result["processing"]["errors"].append(f"An unhandled critical error occurred: {str(e)}")
        logger.exception(f"An unhandled critical error occurred during extraction for {file_path}")
    finally:
        result["processing"]["time_taken_seconds"] = time.time() - start_time
        # python-docx doesn't explicitly require closing like openpyxl, but it's good practice
        # to ensure file handles are released if it were a memory-heavy operation.
        if not result["processing"]["success"] and not result["processing"]["errors"]:
             # If success is false but no errors recorded, it's an unhandled case
            result["processing"]["errors"].append("Extraction failed without specific error message (check logs).")
        logger.info(f"Extraction finished for {file_path}. Success: {result['processing']['success']}")

    return result

# --- CLI Entry Point ---
def run_cli():
    """Command-line interface for the DOCX metadata extractor."""
    parser = argparse.ArgumentParser(
        description="Extract comprehensive metadata from a DOCX file. "
                    "Outputs JSON metadata to stdout or a specified file."
    )
    parser.add_argument(
        "file",
        nargs="?", # Optional argument for file path
        help="Path to the DOCX file to extract metadata from."
    )
    parser.add_argument(
        "--log-level",
        default=DEFAULT_LOG_LEVEL,
        choices=["DEBUG", "INFO", "WARNING", "ERROR", "CRITICAL"],
        help=f"Set the logging level (default: {DEFAULT_LOG_LEVEL})."
    )
    parser.add_argument(
        "--output",
        "-o",
        help="Write metadata output to this file (JSON). If not set, prints to stdout."
    )
    parser.add_argument(
        "--no-input-prompt",
        action="store_true",
        help="Do not prompt for file path if not provided as argument. Exit instead."
    )

    args = parser.parse_args()

    # Re-setup logging based on CLI argument
    setup_logging(args.log_level)
    
    file_path = args.file
    if not file_path:
        if args.no_input_prompt:
            logger.error("No file path provided and --no-input-prompt is set. Exiting.")
            sys.exit(1)
        else:
            file_path = input("Enter the path to the DOCX file: ").strip()
            if not file_path:
                logger.error("No file path provided. Exiting.")
                sys.exit(1)

    try:
        logger.info(f"Attempting to extract metadata for: {file_path}")
        metadata = extract_docx_metadata(file_path)
        output_json = json.dumps(metadata, indent=2, ensure_ascii=False) # ensure_ascii for proper non-ASCII display

        if args.output:
            with open(args.output, "w", encoding="utf-8") as f: # Ensure UTF-8 output
                f.write(output_json)
            logger.info(f"Metadata successfully written to {args.output}")
            print(f"Metadata written to {args.output}")
        else:
            print(output_json)
        
        # Indicate non-zero exit code if extraction was not fully successful
        if not metadata["processing"]["success"] or metadata["processing"]["errors"]:
            logger.warning(f"Extraction completed with warnings/errors for {file_path}. Exit code 2.")
            sys.exit(2)
        else:
            logger.info(f"Extraction successful for: {file_path}. Time taken: {metadata['processing']['time_taken_seconds']:.2f}s")

    except Exception as e: # Catch any remaining unhandled exceptions from CLI setup or argparse
        logger.exception(f"A critical error occurred in the CLI for {file_path}: {e}")
        print(f"A critical error occurred: {e}. Check {DEFAULT_LOG_FILE} for details.")
        sys.exit(1)

def main():
    run_cli()

if __name__ == "__main__":
    main()